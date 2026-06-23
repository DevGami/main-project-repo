import os
from typing import List, Any
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    CSVLoader,
)
from langchain_core.documents import Document


def _load_docx(file_path: str) -> List[Document]:
    """Load a .docx Word file using python-docx and return LangChain Documents."""
    import docx
    doc = docx.Document(file_path)
    full_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [Document(page_content=full_text, metadata={"source": file_path})]


def _load_json_as_text(file_path: str) -> List[Document]:
    """Load JSON/JSONL as plain text — no jq dependency."""
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    return [Document(page_content=content, metadata={"source": file_path})]


def load_all_documents(data_dir: str) -> List[Any]:
    all_docs = []

    pdf_files, txt_files, csv_files, json_files, docx_files = [], [], [], [], []

    supported_text_exts = {
        ".txt", ".sql", ".yaml", ".yml", ".md", ".html",
        ".htm", ".xml", ".log", ".py", ".js", ".ts", ".css",
    }

    for f in os.listdir(data_dir):
        if f.startswith("."):
            continue
        ext = os.path.splitext(f)[1].lower()
        full = os.path.join(data_dir, f)
        if ext == ".pdf":
            pdf_files.append(full)
        elif ext == ".csv":
            csv_files.append(full)
        elif ext in (".json", ".jsonl"):
            json_files.append(full)
        elif ext in (".docx", ".doc"):
            docx_files.append(full)
        elif ext in supported_text_exts:
            txt_files.append(full)
        else:
            # Attempt as text for unknown extensions
            txt_files.append(full)

    print(f"[DEBUG] PDFs={len(pdf_files)} | TXT/other={len(txt_files)} | CSV={len(csv_files)} | JSON={len(json_files)} | DOCX={len(docx_files)}")

    for path in pdf_files:
        try:
            docs = PyPDFLoader(path).load()
            all_docs.extend(docs)
            print(f"[DEBUG] Loaded {len(docs)} pages from {os.path.basename(path)}")
        except Exception as e:
            print(f"[ERROR] PDF {path}: {e}")

    for path in txt_files:
        try:
            docs = TextLoader(path, encoding="utf-8").load()
            all_docs.extend(docs)
        except UnicodeDecodeError:
            try:
                docs = TextLoader(path, encoding="latin-1").load()
                all_docs.extend(docs)
            except Exception as e:
                print(f"[ERROR] Text {path}: {e}")
        except Exception as e:
            print(f"[ERROR] Text {path}: {e}")

    for path in csv_files:
        try:
            docs = CSVLoader(path).load()
            all_docs.extend(docs)
        except Exception as e:
            print(f"[ERROR] CSV {path}: {e}")

    for path in json_files:
        try:
            docs = _load_json_as_text(path)
            all_docs.extend(docs)
        except Exception as e:
            print(f"[ERROR] JSON {path}: {e}")

    for path in docx_files:
        try:
            docs = _load_docx(path)
            all_docs.extend(docs)
            print(f"[DEBUG] Loaded DOCX {os.path.basename(path)}")
        except Exception as e:
            print(f"[ERROR] DOCX {path}: {e}")

    print(f"[DEBUG] Total documents loaded: {len(all_docs)}")
    return all_docs